"""
rag_utils.py
Essential utility functions for growbro-worker vectorstore creation.
"""

import os
import time
import requests
import mimetypes
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Firecrawl SDK
try:
    from firecrawl import FirecrawlApp, ScrapeOptions
except ImportError:
    FirecrawlApp = None
    ScrapeOptions = None

def extract_file_text(file_path):
    """
    Extract text from a file (PDF, TXT, DOCX, etc.).
    Returns a string or None if extraction fails.
    """
    if not os.path.exists(file_path):
        print(f"[extract_file_text] File not found: {file_path}")
        return None
    mime, _ = mimetypes.guess_type(file_path)
    try:
        if mime and mime.startswith("text"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif mime == "application/pdf" or file_path.lower().endswith(".pdf"):
            try:
                from PyPDF2 import PdfReader
            except ImportError:
                print("[extract_file_text] PyPDF2 not installed. Skipping PDF extraction.")
                return None
            try:
                reader = PdfReader(file_path)
                return "\n".join(page.extract_text() or "" for page in reader.pages)
            except Exception as e:
                print(f"[extract_file_text] PDF extraction failed: {e}")
                return None
        elif file_path.lower().endswith(".docx"):
            try:
                import docx
            except ImportError:
                print("[extract_file_text] python-docx not installed. Skipping DOCX extraction.")
                return None
            try:
                doc = docx.Document(file_path)
                return "\n".join([p.text for p in doc.paragraphs])
            except Exception as e:
                print(f"[extract_file_text] DOCX extraction failed: {e}")
                return None
        else:
            print(f"[extract_file_text] Unsupported file type: {file_path}")
            return None
    except Exception as e:
        print(f"[extract_file_text] Extraction error: {e}")
        return None

def extract_website_text_with_firecrawl(urls, min_words=10, firecrawl_api_key=None, formats=['markdown'], limit=30, return_analytics=False, session_cookie=None, deep_crawl=False):
    """
    Extracts website text using Firecrawl. Falls back to generic_extract_website_text if Firecrawl fails or is unavailable.
    Returns: list of LangChain Document objects
    If return_analytics=True, returns (documents, analytics_dict) where analytics_dict has keys 'pages_crawled', 'urls_crawled'.
    """
    api_key = firecrawl_api_key or os.environ.get("FIRECRAWL_API_KEY")
    all_documents = []
    urls_crawled = []
    if deep_crawl and len(urls) > 0:
        url = urls[0]
        try:
            # Ensure URL has scheme
            if not url.startswith("http://") and not url.startswith("https://"):
                url = "https://" + url
            print(f"[Firecrawl Debug] session_cookie: {session_cookie}")
            print(f"[Firecrawl Debug] Deep crawl target url: {url}")
            if session_cookie:
                print("[Firecrawl Debug] Path: REST API /v1/crawl with session_cookie")
                # Use Firecrawl REST API with session cookie
                headers = {
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json",
                    "Cookie": session_cookie
                }
                data = {
                    "url": url,
                    "formats": formats,
                    "limit": limit,
                    "crawlEntireDomain": True  # Enable crawling of entire domain to find missed links
                    # No maxDepth parameter for deep crawl
                }
                print("[Firecrawl Debug] Using /v1/crawl endpoint for deep crawl")
                resp = requests.post("https://api.firecrawl.dev/v1/crawl", headers=headers, json=data, timeout=60)
                resp.raise_for_status()
                result = resp.json()
                for item in result.get('data', []):
                    content = item.get('markdown') or item.get('html') or ""
                    metadata = item.get('metadata', {})
                    print(f"[Firecrawl Debug] Raw metadata from API: {metadata}")
                    # Ensure 'source' is always set for vector deletion
                    if 'source' not in metadata:
                        computed_source = metadata.get('sourceURL') or metadata.get('url') or url
                        print(f"[Firecrawl Debug] Computed 'source' for Document: {computed_source}")
                        metadata['source'] = computed_source
                    print(f"[Firecrawl Debug] Final metadata for Document: {metadata}")
                    all_documents.append(Document(page_content=content, metadata=metadata))
                    page_url = metadata.get('url') or url
                    urls_crawled.append(page_url)
            else:
                print("[Firecrawl Debug] Path: SDK app.crawl_url (no session_cookie)")
                if FirecrawlApp is None or ScrapeOptions is None:
                    print("[Firecrawl Debug] SDK unavailable, using REST /v1/crawl instead")
                    headers = {
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json"
                    }
                    data = {
                        "url": url,
                        "formats": formats,
                        "limit": limit,
                        "crawlEntireDomain": True
                    }
                    try:
                        resp = requests.post("https://api.firecrawl.dev/v1/crawl", headers=headers, json=data, timeout=60)
                        resp.raise_for_status()
                    except requests.exceptions.HTTPError as http_err:
                        error_text = http_err.response.text if http_err.response is not None else ""
                        print(f"[Firecrawl Debug] REST deep crawl HTTPError: {http_err} | Response: {error_text}")
                        raise
                    result = resp.json()
                    for item in result.get('data', []):
                        content = item.get('markdown') or item.get('html') or ""
                        metadata = item.get('metadata', {})
                        if 'source' not in metadata:
                            computed_source = metadata.get('sourceURL') or metadata.get('url') or url
                            print(f"[Firecrawl Debug] Computed 'source' for REST deep crawl Document: {computed_source}")
                            metadata['source'] = computed_source
                        print(f"[Firecrawl Debug] Final metadata for REST deep crawl Document: {metadata}")
                        all_documents.append(Document(page_content=content, metadata=metadata))
                        page_url = metadata.get('url') or url
                        urls_crawled.append(page_url)
                else:
                    # Use Firecrawl SDK for public crawling
                    app = FirecrawlApp(api_key=api_key)
                    crawl_result = app.crawl_url(url, limit=limit, scrape_options=ScrapeOptions(formats=formats))
                    print(f"[Firecrawl Debug] Full crawl_result: {crawl_result}")
                    if hasattr(crawl_result, 'status') and crawl_result.status == 'completed':
                        status = crawl_result
                    else:
                        crawl_id = crawl_result.id
                        while True:
                            status = app.check_crawl_status(crawl_id)
                            print(f"[Firecrawl Debug] Polling crawl status: {status}")
                            if status.status == 'completed':
                                break
                            elif status.status == 'failed':
                                raise RuntimeError(f"Firecrawl crawl failed: {status}")
                            time.sleep(3)
                    print(f"[Firecrawl Debug] Final crawl status: {status}")
                    for item in status.data:
                        content = getattr(item, 'markdown', None) or getattr(item, 'html', None) or ""
                        metadata = getattr(item, 'metadata', {})
                        # Ensure 'source' is always set for vector deletion
                        if 'source' not in metadata:
                            computed_source = metadata.get('sourceURL') or metadata.get('url') or url
                            print(f"[Firecrawl Debug] Computed 'source' for deep SDK Document: {computed_source}")
                            metadata['source'] = computed_source
                        print(f"[Firecrawl Debug] Final metadata for deep SDK Document: {metadata}")
                        all_documents.append(Document(page_content=content, metadata=metadata))
                        page_url = metadata.get('url') or url
                        urls_crawled.append(page_url)
        except Exception as e:
            print(f"[Firecrawl] Error deep crawling {url}: {e}. Please Contact Growbro")
    elif not deep_crawl:
        for url in urls:
            try:
                # Ensure URL has scheme
                if not url.startswith("http://") and not url.startswith("https://"):
                    url = "https://" + url
                if session_cookie:
                    # Use Firecrawl REST API with session cookie
                    headers = {
                        "Authorization": f"Bearer {api_key}",
                        "Content-Type": "application/json",
                        "Cookie": session_cookie
                    }
                    data = {
                        "url": url,
                        "formats": formats,
                        "timeout": 60000
                    }
                    resp = requests.post("https://api.firecrawl.dev/v1/scrape", headers=headers, json=data, timeout=90)
                    resp.raise_for_status()
                    result = resp.json()
                    # v1 /scrape returns data as a single object, not an array
                    item = result.get('data', {})
                    if item:
                        content = item.get('markdown') or item.get('html') or ""
                        metadata = item.get('metadata', {})
                        # Ensure 'source' is always set for vector deletion
                        if 'source' not in metadata:
                            computed_source = metadata.get('sourceURL') or metadata.get('url') or url
                            print(f"[Firecrawl Debug] Computed 'source' for non-deep REST API Document: {computed_source}")
                            metadata['source'] = computed_source
                        print(f"[Firecrawl Debug] Final metadata for non-deep REST API Document: {metadata}")
                        all_documents.append(Document(page_content=content, metadata=metadata))
                        page_url = metadata.get('url') or url
                        urls_crawled.append(page_url)
                else:
                    if FirecrawlApp is None or ScrapeOptions is None:
                        print("[Firecrawl Debug] SDK unavailable, using REST /v1/scrape instead")
                        headers = {
                            "Authorization": f"Bearer {api_key}",
                            "Content-Type": "application/json"
                        }
                        data = {
                            "url": url,
                            "formats": formats,
                            "timeout": 60000
                        }
                        try:
                            resp = requests.post("https://api.firecrawl.dev/v1/scrape", headers=headers, json=data, timeout=90)
                            resp.raise_for_status()
                        except requests.exceptions.HTTPError as http_err:
                            error_text = http_err.response.text if http_err.response is not None else ""
                            print(f"[Firecrawl Debug] REST non-deep HTTPError: {http_err} | Response: {error_text}")
                            raise
                        result = resp.json()
                        # v1 /scrape returns data as a single object, not an array
                        item = result.get('data', {})
                        if item:
                            content = item.get('markdown') or item.get('html') or ""
                            metadata = item.get('metadata', {})
                            if 'source' not in metadata:
                                computed_source = metadata.get('sourceURL') or metadata.get('url') or url
                                print(f"[Firecrawl Debug] Computed 'source' for REST non-deep Document: {computed_source}")
                                metadata['source'] = computed_source
                            print(f"[Firecrawl Debug] Final metadata for REST non-deep Document: {metadata}")
                            all_documents.append(Document(page_content=content, metadata=metadata))
                            page_url = metadata.get('url') or url
                            urls_crawled.append(page_url)
                    else:
                        # Use Firecrawl SDK for public crawling
                        app = FirecrawlApp(api_key=api_key)
                        crawl_result = app.crawl_url(url, limit=1, scrape_options=ScrapeOptions(formats=formats))  # Only crawl this page, do not follow links
                        if hasattr(crawl_result, 'status') and crawl_result.status == 'completed':
                            status = crawl_result
                        else:
                            crawl_id = crawl_result.id
                            while True:
                                status = app.check_crawl_status(crawl_id)
                                if status.status == 'completed':
                                    break
                                elif status.status == 'failed':
                                    raise RuntimeError(f"Firecrawl crawl failed: {status}")
                                time.sleep(3)
                        for item in status.data:
                            content = getattr(item, 'markdown', None) or getattr(item, 'html', None) or ""
                            metadata = getattr(item, 'metadata', {})
                            # Ensure 'source' is always set for vector deletion
                            if 'source' not in metadata:
                                computed_source = metadata.get('sourceURL') or metadata.get('url') or url
                                print(f"[Firecrawl Debug] Computed 'source' for non-deep SDK Document: {computed_source}")
                                metadata['source'] = computed_source
                            print(f"[Firecrawl Debug] Final metadata for non-deep SDK Document: {metadata}")
                            all_documents.append(Document(page_content=content, metadata=metadata))
                            page_url = metadata.get('url') or url
                            urls_crawled.append(page_url)
            except Exception as e:
                print(f"[Firecrawl] Error crawling {url}: {e}. Please Contact Growbro")
                # Fallback for this URL only
                # Optionally could call generic_extract_website_text here and add to docs/analytics
    if return_analytics:
        # Remove duplicates and nulls
        urls_crawled = [u for u in set(urls_crawled) if u]
        return all_documents, {"pages_crawled": len(urls_crawled), "urls_crawled": urls_crawled}
    return all_documents

def get_text_splitter():
    """
    Returns a RecursiveCharacterTextSplitter for document chunking.
    """
    return RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        separators=["\n\n", "\n", ".", "!", "?", ",", " ", ""]
    )

def aggregate_crawl_analytics(website_analytics, link_analytics, file_analytics):
    """
    Aggregate analytics from website crawling, link processing, and file processing.
    """
    analytics = {
        "pages_crawled": 0,
        "urls_crawled": [],
        "files_indexed": 0
    }
    
    if website_analytics:
        analytics["pages_crawled"] += website_analytics.get("pages_crawled", 0)
        analytics["urls_crawled"].extend(website_analytics.get("urls_crawled", []))
    
    if link_analytics:
        analytics["pages_crawled"] += link_analytics.get("pages_crawled", 0)
        analytics["urls_crawled"].extend(link_analytics.get("urls_crawled", []))
    
    if file_analytics:
        analytics["files_indexed"] = file_analytics.get("files_indexed", 0)
    
    return analytics